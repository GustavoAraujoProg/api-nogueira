from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import base64
from io import BytesIO

def gerar_docx_documento(dados_documento, conteudo, dados_escritorio=None):
    document = Document()
    for section in document.sections:
        section.top_margin = Cm(3); section.bottom_margin = Cm(2)
        section.left_margin = Cm(3); section.right_margin = Cm(2)
    
    # 1. CABEÇALHO
    if dados_escritorio and dados_escritorio.get('logoBase64'):
        try:
            logo_data = dados_escritorio['logoBase64'].split(',')[1] if ',' in dados_escritorio['logoBase64'] else dados_escritorio['logoBase64']
            p_logo = document.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(BytesIO(base64.b64decode(logo_data)), width=Cm(2.5))
            document.add_paragraph().paragraph_format.space_after = Pt(6)
        except: pass

    # 2. CONTEÚDO
    conteudo = conteudo.replace('<br>', '\n').replace('<br/>', '\n').replace('</div>', '\n')
    partes = re.split(r'(<div[^>]*jurisprudencia-citacao[^>]*>[\s\S]*?)', conteudo)
    
    for parte in partes:
        if 'jurisprudencia-citacao' in parte:
            texto_limpo = re.sub(r'<[^>]+>', '', re.sub(r'<div[^>]*>', '', parte)).strip().replace('"', '')
            if texto_limpo:
                document.add_paragraph()
                p_jur = document.add_paragraph()
                p_jur.paragraph_format.left_indent = Cm(4); p_jur.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p_jur.add_run(f'"{texto_limpo}"'); run.italic = True; run.font.name = 'Times New Roman'; run.font.size = Pt(10)
                document.add_paragraph()
        else:
            texto_puro = re.sub(r'<[^>]+>', '', parte)
            for linha in texto_puro.split('\n'):
                linha = linha.strip()
                if not linha: continue
                p = document.add_paragraph()
                if (linha.isupper() and 5 < len(linha) < 100) or re.match(r'^[IVX]+\s*-', linha):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if linha.isupper() and not re.match(r'^[IVX]', linha): p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(linha); run.bold = True
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent = Cm(2)
                    run = p.add_run(linha)
                run.font.name = 'Times New Roman'; run.font.size = Pt(12)

    # 3. RODAPÉ
    document.add_paragraph(); document.add_paragraph()
    p_div = document.add_paragraph("_"*50); p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.runs[0].font.color.rgb = RGBColor(150, 150, 150); p_div.runs[0].font.size = Pt(8)
    
    if dados_escritorio:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(dados_escritorio.get('nome', '').upper())
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Arial'; r.font.color.rgb = RGBColor(197, 160, 89)
        
        info = f"{dados_escritorio.get('endereco', '')} • {dados_escritorio.get('telefone', '')}\n{dados_escritorio.get('email', '')}"
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(info); r2.font.size = Pt(8); r2.font.name = 'Arial'; r2.font.color.rgb = RGBColor(80, 80, 80)

    out = io.BytesIO(); document.save(out); out.seek(0)
    return out.getvalue()